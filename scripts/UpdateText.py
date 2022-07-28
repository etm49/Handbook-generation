import os, sys

import docx

import pandas as pd
import json
import re

from enums import stats, outputs
from commons import data_path, input_path, output_path
from utlis import get_inputs

# Define statistic calculating function

def statCalculator(code_list):
	###
	#  Calculates the statical values definied in enums.py based on the code_list. E.g.: if {ohi-index,minSids, countryName}, this function returns name of the SIDS country with the minimum value for ohi-index for the given year 
	# ###

    assert len(code_list) in [1,3], "code is not properly desigined, should flow format {ohi-index,minSids, countryName} "
    
    # Return year
    if code_list[-1] == outputs.year.name:
        return str(year)
    else:
        requiredStat = code_list[1]
        requiredOutput = code_list[2]
        indicator = code_list[0]
        # Calculate stat
        if (requiredStat == stats.minSids.name):
            data = indicatorData[(indicatorData["Indicator Code"]==indicator)&(indicatorData["Country Code"].isin(SIDS))][[str(year),"Country Code"]].min()
        elif (requiredStat == stats.minGlobal.name):
            data = indicatorData[indicatorData["Indicator Code"]==indicator][[str(year),"Country Code"]].min()
        elif (requiredStat == stats.maxSids.name):
            data = indicatorData[(indicatorData["Indicator Code"]==indicator)&(indicatorData["Country Code"].isin(SIDS))][[str(year),"Country Code"]].max()
        elif (requiredStat == stats.maxGlobal.name):
            data = indicatorData[indicatorData["Indicator Code"]==indicator][[str(year),"Country Code"]].max()
        elif (requiredStat == stats.aveSids.name):
            data = indicatorData[(indicatorData["Indicator Code"]==indicator)&(indicatorData["Country Code"].isin(SIDS))][[str(year),"Country Code"]].mean()
        elif (requiredStat == stats.aveGlobal.name):
            data = indicatorData[indicatorData["Indicator Code"]==indicator][[str(year),"Country Code"]].mean()
        elif (requiredStat == stats.minPacific.name):
            data = indicatorData[(indicatorData["Indicator Code"]==indicator)&(indicatorData["Country Code"].isin(pacificRegion))][[str(year),"Country Code"]].min()
        elif (requiredStat == stats.maxPacific.name):
            data = indicatorData[(indicatorData["Indicator Code"]==indicator)&(indicatorData["Country Code"].isin(pacificRegion))][[str(year),"Country Code"]].max()
        elif (requiredStat == stats.avePacific.name):
            data = indicatorData[(indicatorData["Indicator Code"]==indicator)&(indicatorData["Country Code"].isin(pacificRegion))][[str(year),"Country Code"]].mean()
        elif (requiredStat == stats.minCarribean.name):
            data = indicatorData[(indicatorData["Indicator Code"]==indicator)&(indicatorData["Country Code"].isin(carribeanRegion))][[str(year),"Country Code"]].min()
        elif (requiredStat == stats.maxCarribean.name):
            data = indicatorData[(indicatorData["Indicator Code"]==indicator)&(indicatorData["Country Code"].isin(carribeanRegion))][[str(year),"Country Code"]].max()
        elif (requiredStat == stats.aveCarribean.name):
            data = indicatorData[(indicatorData["Indicator Code"]==indicator)&(indicatorData["Country Code"].isin(carribeanRegion))][[str(year),"Country Code"]].mean()
        elif (requiredStat == stats.minAIS.name):
            data = indicatorData[(indicatorData["Indicator Code"]==indicator)&(indicatorData["Country Code"].isin(aisRegion))][[str(year),"Country Code"]].min()
        elif (requiredStat == stats.maxAIS.name):
            data = indicatorData[(indicatorData["Indicator Code"]==indicator)&(indicatorData["Country Code"].isin(aisRegion))][[str(year),"Country Code"]].max()
        elif (requiredStat == stats.aveAIS.name):
            data = indicatorData[(indicatorData["Indicator Code"]==indicator)&(indicatorData["Country Code"].isin(aisRegion))][[str(year),"Country Code"]].mean()
        
        # Return required output
        if requiredOutput == outputs.value.name:
            return str(round(data[str(year)],2))
        elif requiredOutput == outputs.countryName.name:
            return sidsNames[data["Country Code"]]
def processor(code):
    code_list = code.strip("{}").split(",")

    return statCalculator(code_list=code_list)


if __name__ == "__main__":

	year = get_inputs("please input year for analysis e.g. 2010")
	year = int(year)

	docName = get_inputs("please input name of docx to be updated from the input files in word doc folder")
	doc = docx.Document(input_path +  docName)




	# Import indicator data for calculating values
	indicatorData = pd.read_csv(data_path+"indicatorData.csv").sort_index(axis=1)
	indicatorMeta = pd.read_csv(data_path+"indicatorMeta.csv")


	# Opening SIDS JSON file
	with open(data_path+'undpSids.json', 'r') as openfile:
		
		# Reading from json file
		undpSids = json.load(openfile)

	with open(data_path+'region.json', 'r') as openfile:
  
		# Reading from json file
		undpSidsRegion = json.load(openfile)


	# Create a list of SIDS iso-3 codes
	SIDS = [] 
	for i in undpSids.keys():
		SIDS.append(undpSids[i])
		
	# create a dictionary with SIDS iso-3 codes as keys and country names as values
	sidsNames = {}
	sidsNames = dict([(value, key) for key, value in undpSids.items()])

	# Create a list of SIDS iso-3 codes by region
	pacificRegion = []
	carribeanRegion = []
	aisRegion = []
	for i in undpSidsRegion:
		if i["region"] == "Caribbean":
			carribeanRegion.append(i["iso"])
		elif i["region"] == "AIS":
			aisRegion.append(i["iso"])
		elif i["region"] == "Pacific":
			pacificRegion.append(i["iso"])
    


	### Replace coded texts in paragraphs (those between curly brackets according to the enumeration classes)
	for p in doc.paragraphs:
		temp = re.findall(r'\{.*?\}', p.text)
		if len(temp) > 0:
			result = [processor(i) for i in temp]
			for ind,val in enumerate(temp):
				p.text = p.text.replace(val, result[ind],1)

	### Replace coded texts in tables (those between curly brackets according to the enumeration classes)
	for table in doc.tables:
		for row in table.rows:
			for cell in row.cells:
				for p in cell.paragraphs:
		
					temp = re.findall(r'\{.*?\}', p.text)
					if len(temp) > 0:
						result = [processor(i) for i in temp]
						for ind,val in enumerate(temp):
							p.text = p.text.replace(val, result[ind],1)

	# save new word document with replaced values
	doc.save(output_path + str(year) +" GENERATED " + docName)

